import streamlit as st
import os
import tempfile
import shutil
import gc
import time
import json
import re
from io import BytesIO
from datetime import datetime

# --- BIBLIOTECAS ---
import whisper
import pandas as pd
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
import pytesseract
from PIL import Image, ImageOps
import fitz  # PyMuPDF
import numpy as np

# --- 1. CONFIGURAÇÃO ---
CAMINHO_TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

if os.path.exists(CAMINHO_TESSERACT):
    pytesseract.pytesseract.tesseract_cmd = CAMINHO_TESSERACT
    STATUS_OCR = True
else:
    STATUS_OCR = False

PASTA_MEMORIA = "./banco_de_dados_nemesis"
ARQUIVO_CONFIG = os.path.join(PASTA_MEMORIA, "nemesis_config.json")
MODELO_ATUAL = "llama3.1"

st.set_page_config(page_title="NEMESIS AI PRO", page_icon="⚖️", layout="wide")

# --- 2. INICIALIZAÇÃO ---
if "vectorstore" not in st.session_state: st.session_state.vectorstore = None
if "caso_selecionado" not in st.session_state: st.session_state.caso_selecionado = None
if "memoria_imediata" not in st.session_state: st.session_state.memoria_imediata = ""
if "messages" not in st.session_state: st.session_state.messages = []
if "ultimas_fontes" not in st.session_state: st.session_state.ultimas_fontes = []

# --- 3. CACHE ---
@st.cache_resource
def get_whisper_model():
    return whisper.load_model("base")

@st.cache_resource
def get_llm():
    return ChatOllama(model=MODELO_ATUAL, temperature=0.0)

@st.cache_resource
def get_embedding_function():
    return OllamaEmbeddings(model="all-minilm")

# --- 4. UTILITÁRIOS ---
def gerar_word(texto):
    doc = DocxDocument()
    doc.add_heading('Relatório Nemesis AI', 0)
    doc.add_paragraph(f"Gerado: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(texto)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def transcrever_audio_whisper(caminho_audio):
    model = get_whisper_model()
    try:
        result = model.transcribe(caminho_audio)
        return result["text"]
    except Exception as e: return f"Erro áudio: {e}"

def ler_planilha(caminho_arquivo):
    try:
        if caminho_arquivo.endswith('.csv'): df = pd.read_csv(caminho_arquivo)
        else: df = pd.read_excel(caminho_arquivo)
        # Converte para markdown para a IA entender a estrutura
        return df.to_markdown(index=False)
    except ImportError:
        return "ERRO CRÍTICO: Instale 'tabulate' (pip install tabulate)"
    except Exception as e:
        return f"Erro planilha: {e}"

# --- 5. CSS ---
st.markdown("""
<style>
    .stApp { background-color: #1e1e1e; color: #e0e0e0; }
    [data-testid="stSidebar"] { background-color: #121212; border-right: 1px solid #333; }
    .stButton button { text-align: left; border: none; background: transparent; color: #cfcfcf; padding: 8px 0px; width: 100%; }
    .stButton button:hover { color: #fff; padding-left: 5px; }
    [data-testid="stPopover"] > button { border: none; background: transparent; color: #a8c7fa; font-weight: bold; font-size: 1.5em; padding: 0px; width: 30px; }
    .source-box { background-color: #25262b; padding: 12px; border-radius: 8px; border-left: 3px solid #6c5ce7; margin-top: 8px; font-size: 0.85em; color: #b2bec3; }
    .raw-text { font-family: monospace; font-size: 0.8em; color: #a29bfe; background-color: #151515; padding: 10px; border-radius: 5px; max-height: 200px; overflow-y: auto; }
</style>
""", unsafe_allow_html=True)

# --- 6. GESTÃO ---
def carregar_config():
    if not os.path.exists(PASTA_MEMORIA): os.makedirs(PASTA_MEMORIA)
    if not os.path.exists(ARQUIVO_CONFIG):
        with open(ARQUIVO_CONFIG, 'w') as f: json.dump({"pinned": [], "trash": []}, f)
    try:
        with open(ARQUIVO_CONFIG, 'r') as f:
            cfg = json.load(f)
            if "trash" not in cfg: cfg["trash"] = []
            return cfg
    except: return {"pinned": [], "trash": []}

def salvar_config(config):
    with open(ARQUIVO_CONFIG, 'w') as f: json.dump(config, f)

def faxina_inicial():
    cfg = carregar_config()
    lixo = cfg.get("trash", [])
    if not lixo: return
    restante = []
    for caso in lixo:
        caminho = os.path.join(PASTA_MEMORIA, caso)
        if os.path.exists(caminho):
            try: shutil.rmtree(caminho) 
            except: restante.append(caso) 
    if len(restante) != len(lixo):
        cfg["trash"] = restante
        salvar_config(cfg)

faxina_inicial()

def acao_excluir(nome_caso):
    st.session_state.vectorstore = None
    st.session_state.caso_selecionado = None
    st.session_state.memoria_imediata = ""
    st.session_state.messages = []
    gc.collect() 
    cfg = carregar_config()
    if nome_caso not in cfg["trash"]: cfg["trash"].append(nome_caso)
    if nome_caso in cfg.get("pinned", []): cfg["pinned"].remove(nome_caso)
    salvar_config(cfg)
    try: shutil.rmtree(os.path.join(PASTA_MEMORIA, nome_caso))
    except: pass 
    st.toast("🗑️ Enviado para lixeira.")
    time.sleep(0.5)
    st.rerun()

def acao_fixar(nome_caso):
    cfg = carregar_config()
    if nome_caso in cfg["pinned"]:
        cfg["pinned"].remove(nome_caso)
        st.toast("Desfixado")
    else:
        cfg["pinned"].append(nome_caso)
        st.toast("📌 Fixado")
    salvar_config(cfg)
    st.rerun()

def acao_renomear(antigo, novo):
    if not novo: return
    novo_limpo = re.sub(r'[^a-zA-Z0-9_-]', '', novo.strip().replace(" ", "_")).strip("_-")
    if not novo_limpo: return
    st.session_state.vectorstore = None
    gc.collect()
    try:
        os.rename(os.path.join(PASTA_MEMORIA, antigo), os.path.join(PASTA_MEMORIA, novo_limpo))
        cfg = carregar_config()
        if antigo in cfg["pinned"]:
            cfg["pinned"].remove(antigo)
            cfg["pinned"].append(novo_limpo)
        salvar_config(cfg)
        st.session_state.caso_selecionado = novo_limpo
        st.toast("✅ Renomeado!")
        st.rerun()
    except: st.error("⚠️ Erro ao renomear.")

def listar_casos_visiveis():
    if not os.path.exists(PASTA_MEMORIA): os.makedirs(PASTA_MEMORIA)
    cfg = carregar_config()
    todos = [f for f in os.listdir(PASTA_MEMORIA) 
             if os.path.isdir(os.path.join(PASTA_MEMORIA, f)) 
             and not f.startswith("_") 
             and f not in cfg.get("trash", [])]
    pinned = [c for c in todos if c in cfg.get("pinned", [])]
    others = [c for c in todos if c not in cfg.get("pinned", [])]
    return sorted(pinned) + sorted(others)

def carregar_banco(nome_caso):
    caminho = os.path.join(PASTA_MEMORIA, nome_caso)
    if not os.path.exists(caminho): os.makedirs(caminho)
    if st.session_state.vectorstore:
        try: st.session_state.vectorstore._client._system.stop()
        except: pass
    return Chroma(collection_name=nome_caso, embedding_function=get_embedding_function(), persist_directory=caminho)

# --- 7. PROCESSAMENTO (COM CORREÇÃO DE LOTE) ---
def processar_arquivos(vectorstore, arquivos):
    if not arquivos: return 0, ""
    docs_acumulados = []
    texto_sessao = ""
    progresso = st.progress(0, text="Processando...")
    
    for i, arquivo in enumerate(arquivos):
        ext = arquivo.name.split('.')[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp.write(arquivo.getvalue())
            tmp_path = tmp.name
        
        texto_extraido = ""
        origem = ""
        
        try:
            if ext in ['jpg', 'jpeg', 'png']:
                origem = "IMAGEM"
                img = Image.open(tmp_path)
                if np.mean(np.array(img.convert("L"))) < 127: img = ImageOps.invert(img.convert("L"))
                texto_extraido = pytesseract.image_to_string(img, lang='por+eng')
            elif ext == 'pdf':
                origem = "PDF"
                doc = fitz.open(tmp_path)
                for pag in doc:
                    t = pag.get_text()
                    if len(t.strip()) < 5: 
                        pix = pag.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        if np.mean(np.array(img.convert("L"))) < 127: img = ImageOps.invert(img.convert("L"))
                        t = pytesseract.image_to_string(img, lang='por+eng')
                    texto_extraido += t + "\n"
                doc.close()
            elif ext in ['wav', 'mp3']:
                origem = "ÁUDIO"
                texto_extraido = transcrever_audio_whisper(tmp_path)
            elif ext in ['xlsx', 'xls', 'csv']:
                origem = "PLANILHA"
                texto_extraido = ler_planilha(tmp_path)

            if texto_extraido and len(texto_extraido.strip()) > 2:
                doc_final = f"--- {origem}: {arquivo.name} ---\n{texto_extraido}\n"
                docs_acumulados.append(Document(page_content=doc_final, metadata={"source_name": arquivo.name}))
                texto_sessao += doc_final + "\n\n"
            else:
                st.warning(f"⚠️ {arquivo.name} vazio.")

        except Exception as e:
            st.error(f"Erro em {arquivo.name}: {e}")
        finally:
            if os.path.exists(tmp_path): os.remove(tmp_path)
        progresso.progress((i + 1) / len(arquivos))
        
    if docs_acumulados and vectorstore:
        progresso.text("Indexando memória (isso pode demorar)...")
        splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
        todos_fragmentos = splitter.split_documents(docs_acumulados)
        
        # --- A CORREÇÃO DO ERRO DE BATCH SIZE AQUI ---
        # ChromaDB tem limite de ~5461. Vamos inserir em lotes de 4000.
        TAMANHO_LOTE = 4000
        total_docs = len(todos_fragmentos)
        
        for i in range(0, total_docs, TAMANHO_LOTE):
            lote = todos_fragmentos[i : i + TAMANHO_LOTE]
            vectorstore.add_documents(lote)
            time.sleep(0.1) # Respiro para o banco
    
    progresso.empty()
    return len(docs_acumulados), texto_sessao

# --- 8. CHAT ---
def fluxo_de_resposta(vectorstore, pergunta):
    imediato = st.session_state.get("memoria_imediata", "")
    ctx_bd = ""
    historico = []
    
    if vectorstore:
        try:
            historico = vectorstore.as_retriever(search_kwargs={"k": 5}).invoke(pergunta)
            ctx_bd = "\n\n".join([d.page_content for d in historico])
        except: pass
    
    st.session_state.ultimas_fontes = historico

    if not imediato and not ctx_bd:
        yield "Sem dados. Anexe um arquivo."
        return

    contexto_final = f"""
    VOCÊ É UM ANALISTA DE DADOS (JURÍDICO/FINANCEIRO).
    DADOS BRUTOS (OCR/PLANILHAS):
    {imediato}
    HISTÓRICO:
    {ctx_bd}
    INSTRUÇÕES:
    1. Responda com base nos dados acima.
    2. NÃO recuse analisar imagens/planilhas. O texto acima é o conteúdo delas.
    """
    
    chain = ChatPromptTemplate.from_template("Analise e responda:\n{question}") | get_llm()
    # Concatena prompt e pergunta para garantir contexto
    full_q = f"{contexto_final}\n\nPERGUNTA: {pergunta}"
    
    for chunk in chain.stream({"question": full_q}):
        yield chunk.content

# --- MAIN ---
def main():
    with st.sidebar:
        st.header("🗂️ Histórico")
        if st.button("➕ Novo Caso", use_container_width=True):
            st.session_state.caso_selecionado = None
            st.session_state.memoria_imediata = ""
            st.session_state.messages = []
            st.rerun()
        st.markdown("---")
        
        casos = listar_casos_visiveis()
        cfg = carregar_config()
        
        for caso in casos:
            c1, c2 = st.columns([0.85, 0.15])
            with c1:
                icone = "📌" if caso in cfg["pinned"] else "📁"
                if st.button(f"{icone} {caso}", key=f"btn_{caso}", use_container_width=True):
                    st.session_state.caso_selecionado = caso
                    st.session_state.memoria_imediata = ""
                    st.session_state.messages = []
                    st.session_state.vectorstore = None 
                    st.rerun()
            with c2:
                with st.popover("⋮"):
                    st.caption(f"Caso: {caso}")
                    if st.button(f"📌 Fixar", key=f"pin_{caso}", use_container_width=True): acao_fixar(caso)
                    nn = st.text_input("Renomear:", value=caso, key=f"inp_{caso}")
                    if st.button("✏️ Salvar", key=f"ren_{caso}", use_container_width=True): acao_renomear(caso, nn)
                    st.divider()
                    if st.button("🗑️ Excluir", key=f"del_{caso}", type="primary", use_container_width=True): acao_excluir(caso)

        if not STATUS_OCR: st.error("🚨 TESSERACT OFF")

    if st.session_state.get("caso_selecionado"):
        st.title(f"⚖️ {st.session_state.caso_selecionado}")
        if st.session_state.vectorstore is None:
            st.session_state.vectorstore = carregar_banco(st.session_state.caso_selecionado)
    else:
        st.markdown("# 👋 Nemesis AI")
        novo = st.text_input("Novo Cliente:", placeholder="Ex: Silva")
        if novo:
            safe = re.sub(r'[^a-zA-Z0-9_-]', '', novo.strip().replace(" ", "_")).strip("_-")
            if not safe: safe="novo"
            st.session_state.caso_selecionado = safe
            st.rerun()

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant":
                doc = gerar_word(msg["content"])
                st.download_button("📄 Baixar DOCX", doc, f"Nemesis_{int(time.time())}.docx", key=f"d_{msg['content'][:5]}")
            if msg.get("fontes"):
                with st.expander("🔍 Fontes"):
                    for d in msg["fontes"]: st.markdown(f"<div class='source-box'>📄 {d.metadata.get('source_name')}</div>", unsafe_allow_html=True)

    if st.session_state.get("caso_selecionado"):
        if st.session_state.get("memoria_imediata"):
            st.markdown(f"<span class='focus-badge'>🔥 FOCO ATIVO</span>", unsafe_allow_html=True)
            with st.expander("👁️ Ver Dados Brutos"): st.markdown(f"<div class='raw-text'>{st.session_state.memoria_imediata[:2000]}</div>", unsafe_allow_html=True)

        with st.expander("📎 Anexar", expanded=False):
            c1, c2 = st.columns([5, 1])
            files = c1.file_uploader("Upload", type=["pdf", "jpg", "png", "xlsx", "csv", "wav"], accept_multiple_files=True, label_visibility="collapsed")
            if c2.button("Processar", use_container_width=True) and files:
                qtd, texto = processar_arquivos(st.session_state.vectorstore, files)
                st.session_state.memoria_imediata = texto
                st.rerun()

        if prompt := st.chat_input("Pergunte..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                resp = st.write_stream(fluxo_de_resposta(st.session_state.vectorstore, prompt))
                ft = st.session_state.get("ultimas_fontes", [])
                st.session_state.messages.append({"role": "assistant", "content": resp, "fontes": ft})
                st.rerun()

if __name__ == "__main__":
    main()